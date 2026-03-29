#!/usr/bin/env python3
"""
generate_wi_docs.py  —  Template-driven WI document generator.
Reads WI_Template.docx for layout, fills placeholders with Excel data.
pip install openpyxl python-docx
"""

import argparse, copy, sys
from pathlib import Path

try:
    from openpyxl import load_workbook
except ImportError:
    sys.exit("ERROR: Run:  pip install openpyxl")
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree
except ImportError:
    sys.exit("ERROR: Run:  pip install python-docx")

SCRIPT_DIR       = Path(__file__).parent
DEFAULT_TEMPLATE = SCRIPT_DIR / "WI_Template.docx"
DEFAULT_OUTPUT   = SCRIPT_DIR / "Output Docs"
EXCEL_FALLBACKS  = ["WI_Database (1).xlsx", "WI_Database.xlsx", "WI Database.xlsx"]

# ── Excel column mappings ─────────────────────────────────────────────────────
HEADER_COLS = {
    "generate?":"Generate",
    "wi_no *":"WI_No","wi_no":"WI_No",
    "document_title *":"Document_Title","document_title":"Document_Title",
    "version *":"Version","version":"Version",
    "effective_date *":"Effective_Date","effective_date":"Effective_Date",
    "review_date":"Review_Date",
    "site_area":"Site_Area",
    "equipment_asset":"Equipment_Asset",
    "prepared_by":"Prepared_By",
    "approved_by":"Approved_By",
    "asset_tag":"Asset_Tag",
    "department":"Department",
    "scope":"Scope",
    "revision_history":"Revision_History",
    "file_name":"File_Name",
    "asset_status":"Asset_Status",
    "purpose":"Purpose",
    "references_definitions":"References_Definitions",
    "notes":"Notes",
    "safety_risk_controls":"Safety_Risk_Controls",
    "prestart_checklist":"PreStart_Checklist",
    "isolation_shutdown_loto_steps":"Isolation_Shutdown_LOTO_Steps",
    "testing_returntoservice_steps":"Testing_ReturnToService_Steps",
}
STEPS_COLS = {
    "wi_no *":"WI_No","wi_no":"WI_No",
    "step_no *":"Step_No","step_no":"Step_No",
    "step_instruction *":"Step_Instruction","step_instruction":"Step_Instruction",
    "step_keypoints_hazards":"Step_KeyPoints_Hazards",
    "step_check":"Step_Check","section":"Section",
}

# ── Placeholder → field mapping ───────────────────────────────────────────────
# Maps what appears in the template to the Excel field key.
# Numbered variants like [Purpose1], [Purpose2] all map to the same field —
# the script handles expansion automatically.
PLACEHOLDER_MAP = {
    "[Document_Title]":    "Document_Title",
    "[WI_No]":             "WI_No",
    "[Version]":           "Version",
    "[Review_Date]":       "Review_Date",
    "[File_Name]":         "File_Name",
    "[Asset_Status]":      "Asset_Status",
    "[Site_Area]":         "Site_Area",
    "[Equipment_Asset]":   "Equipment_Asset",
    "[Asset_Tag]":         "Asset_Tag",
    "[Prepared_By]":       "Prepared_By",
    "[Approved_By]":       "Approved_By",
    "[Department]":        "Department",
    "[Scope]":             "Scope",
    "[Purpose]":           "Purpose",
    "[References_Definitions]": "References_Definitions",
    "[Notes]":             "Notes",
    "[Safety_Risk_Controls]": "Safety_Risk_Controls",
    "[Isolation_Shutdown_LOTO_Steps]": "Isolation_Shutdown_LOTO_Steps",
}

# Fields that can have numbered variants: [Field1], [Field2] etc.
MULTILINE_FIELDS = {
    "Purpose", "References_Definitions", "Notes", "Safety_Risk_Controls",
    "PreStart_Checklist", "Testing_ReturnToService_Steps", "Isolation_Shutdown_LOTO_Steps",
}


# ── XML / cell helpers ────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def set_cell_borders(cell, color="BBBBBB", size="4"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),size)
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),color)
        tcB.append(b)
    tcPr.append(tcB)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top',top),('bottom',bottom),('left',left),('right',right)):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:w'),str(val))
        el.set(qn('w:type'),'dxa'); tcMar.append(el)
    tcPr.append(tcMar)

def set_col_width(cell, w):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'),str(w))
    tcW.set(qn('w:type'),'dxa'); tcPr.append(tcW)

def get_cell_bg(cell):
    shd = cell._tc.find('.//' + qn('w:shd'))
    if shd is not None:
        return shd.get(qn('w:fill'), '').upper()
    return None

def copy_run_format(src_run, dst_run):
    """Copy font properties from one run to another."""
    if src_run.font.name:     dst_run.font.name  = src_run.font.name
    if src_run.font.size:     dst_run.font.size  = src_run.font.size
    if src_run.font.bold is not None: dst_run.font.bold = src_run.font.bold
    if src_run.font.color and src_run.font.color.type:
        dst_run.font.color.rgb = src_run.font.color.rgb


# ── Placeholder replacement in a paragraph ───────────────────────────────────
def get_para_text(para):
    return "".join(r.text for r in para.runs)

def replace_para_text(para, new_text):
    """Replace all run text in a paragraph with new_text, keep first run's format."""
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""

def find_placeholder_in_para(para):
    """Return the placeholder tag found in this paragraph, or None."""
    text = get_para_text(para)
    if '[' in text and ']' in text:
        start = text.index('[')
        end   = text.index(']', start) + 1
        return text[start:end]
    return None


# ── Multi-line cell expansion ─────────────────────────────────────────────────
def fill_multiline_cell(cell, lines):
    """
    Fill a table cell with multiple lines of text.
    Each line becomes a separate paragraph, preserving the cell's existing format.
    """
    if not lines:
        lines = [""]

    # Get format from first existing paragraph
    ref_para = cell.paragraphs[0]
    ref_runs = ref_para.runs

    # Clear all existing paragraphs
    tc = cell._tc
    for p in tc.findall(qn('w:p')):
        tc.remove(p)

    for i, line in enumerate(lines):
        p_el = OxmlElement('w:p')
        tc.append(p_el)

        # Copy paragraph properties from reference if available
        if ref_para._p.find(qn('w:pPr')) is not None:
            new_pPr = copy.deepcopy(ref_para._p.find(qn('w:pPr')))
            p_el.append(new_pPr)

        if line.strip():
            r_el = OxmlElement('w:r')
            p_el.append(r_el)

            # Copy run properties
            if ref_runs:
                src_rPr = ref_runs[0]._r.find(qn('w:rPr'))
                if src_rPr is not None:
                    r_el.append(copy.deepcopy(src_rPr))

            t_el = OxmlElement('w:t')
            t_el.text = line
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r_el.append(t_el)


def expand_numbered_placeholders(cell, field_name, lines):
    """
    Replace [FieldName1], [FieldName2] ... numbered placeholders in a cell
    with the actual lines of text. Extra template rows beyond available data
    are cleared; if there are more lines than template rows, extra lines
    are appended to the last placeholder paragraph.
    """
    paras = cell.paragraphs
    placeholder_paras = []
    for para in paras:
        text = get_para_text(para)
        if f'[{field_name}' in text and ']' in text:
            placeholder_paras.append(para)

    if not placeholder_paras:
        return False  # Nothing to replace

    if not lines:
        lines = [""]

    # Fill each placeholder paragraph with its corresponding line
    for i, para in enumerate(placeholder_paras):
        if i < len(lines):
            replace_para_text(para, lines[i])
        else:
            replace_para_text(para, "")  # Clear unused placeholders

    # If more lines than placeholder slots, append to last para's cell
    if len(lines) > len(placeholder_paras):
        tc = cell._tc
        ref_para = placeholder_paras[-1]
        for extra_line in lines[len(placeholder_paras):]:
            p_el = OxmlElement('w:p')
            # Insert after last placeholder paragraph
            ref_para._p.addnext(p_el)
            ref_para = ref_para  # keep ref

            if ref_para._p.find(qn('w:pPr')) is not None:
                p_el.append(copy.deepcopy(ref_para._p.find(qn('w:pPr'))))

            if extra_line.strip():
                r_el = OxmlElement('w:r')
                p_el.append(r_el)
                if ref_para.runs:
                    src_rPr = ref_para.runs[0]._r.find(qn('w:rPr'))
                    if src_rPr is not None:
                        r_el.append(copy.deepcopy(src_rPr))
                t_el = OxmlElement('w:t')
                t_el.text = extra_line
                t_el.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
                r_el.append(t_el)

    return True


# ── Steps table replacement ───────────────────────────────────────────────────
def find_steps_table(doc, marker):
    """Find a steps table by looking for a specific placeholder pattern."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if marker in cell.text:
                    return table
    return None


def replace_steps_table(table, steps):
    """
    Replace sample rows in a steps table with real data.
    Clones the first data row as the template for new rows.
    """
    if not steps:
        # Remove all data rows, leave header
        for row in list(table.rows)[1:]:
            table._tbl.remove(row._tr)
        return

    # Keep header row, use first data row as template
    data_rows = list(table.rows)[1:]
    if not data_rows:
        return

    template_row = data_rows[0]
    col_widths = []
    for cell in template_row.cells:
        tcW = cell._tc.find('.//' + qn('w:tcW'))
        col_widths.append(int(tcW.get(qn('w:w'), 2000)) if tcW is not None else 2000)

    # Get background colour from first data row
    bg_colors = []
    for cell in template_row.cells:
        shd = cell._tc.find('.//' + qn('w:shd'))
        bg_colors.append(shd.get(qn('w:fill'), 'FFFFFF') if shd else 'FFFFFF')

    ALT_BG = "EBF3FB"

    # Remove all existing data rows
    for row in data_rows:
        table._tbl.remove(row._tr)

    # Build new rows from steps data
    for ri, step in enumerate(steps):
        bg = 'FFFFFF' if ri % 2 == 0 else ALT_BG
        tr = copy.deepcopy(template_row._tr)
        table._tbl.append(tr)

        # Get cells from the new row
        cells = tr.findall('.//' + qn('w:tc'))
        n_cols = len(cells)

        col_data = [str(step.get("Step_No", ri+1))]
        col_data.append(str(step.get("Step_Instruction", "")))
        if n_cols >= 3:
            col_data.append(str(step.get("Step_Check", "☐")))

        for ci, tc_el in enumerate(cells[:len(col_data)]):
            # Update background
            shd = tc_el.find('.//' + qn('w:shd'))
            if shd is not None:
                shd.set(qn('w:fill'), bg)

            # Clear existing paragraphs and set text
            for p in tc_el.findall(qn('w:p')):
                tc_el.remove(p)

            text = col_data[ci] if ci < len(col_data) else ""
            for li, line in enumerate(text.split('\n')):
                p_el = OxmlElement('w:p')
                tc_el.append(p_el)
                if line.strip():
                    r_el = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    sz = OxmlElement('w:sz'); sz.set(qn('w:val'),'18')
                    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'),'18')
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'),'Arial'); rFonts.set(qn('w:hAnsi'),'Arial')
                    rPr.extend([rFonts, sz, szCs]); r_el.append(rPr)
                    t_el = OxmlElement('w:t'); t_el.text = line
                    t_el.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
                    r_el.append(t_el); p_el.append(r_el)


# ── Simple single-value replacement ──────────────────────────────────────────
def replace_simple_placeholders(doc, wi):
    """Replace all single-value [Tag] placeholders everywhere in the document."""
    def sub(text):
        for key, field in PLACEHOLDER_MAP.items():
            text = text.replace(key, str(wi.get(field, "") or ""))
        return text

    def process_para(para):
        full = get_para_text(para)
        new  = sub(full)
        if new != full:
            replace_para_text(para, new)

    def process_element(element):
        for para in element.paragraphs:
            process_para(para)
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_element(cell)

    process_element(doc)
    for section in doc.sections:
        for part in [section.header, section.footer]:
            process_element(part)


# ── Multi-line field expansion in body tables ─────────────────────────────────
def expand_multiline_fields(doc, wi):
    """
    For each table cell containing numbered placeholders like [Purpose1][Purpose2],
    expand them with the actual lines from the wi data field.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                for field in MULTILINE_FIELDS:
                    if f'[{field}' in cell_text:
                        raw_value = wi.get(field, "") or ""
                        lines = [l for l in raw_value.split('\n') if l.strip()]
                        if not lines:
                            lines = [""]
                        expand_numbered_placeholders(cell, field, lines)
                        break  # one field per cell


# ── Main document builder ─────────────────────────────────────────────────────
def build_document(template_path, wi, steps):
    doc = Document(str(template_path))

    # 1. Replace simple single-value placeholders (header, footer, info tables)
    replace_simple_placeholders(doc, wi)

    # 2. Expand multi-line fields in body tables
    expand_multiline_fields(doc, wi)

    # 3. Replace steps tables
    pre_steps  = [s for s in steps if s.get("Section") == "Pre-start"]
    task_steps = [s for s in steps if s.get("Section") == "Task execution"]
    test_steps = [s for s in steps if s.get("Section") == "Testing"]
    loto_steps = [s for s in steps if s.get("Section") == "LOTO"]
    all_steps  = steps if not any([pre_steps, task_steps, test_steps, loto_steps]) else []

    prestart_tbl = find_steps_table(doc, "[PreStart_Checklist")
    task_tbl     = find_steps_table(doc, "[Step_1_Instruction]")
    test_tbl     = find_steps_table(doc, "[Testing_ReturnToService_Steps")

    if prestart_tbl: replace_steps_table(prestart_tbl, pre_steps or loto_steps)
    if task_tbl:     replace_steps_table(task_tbl,     task_steps or all_steps)
    if test_tbl:     replace_steps_table(test_tbl,     test_steps)

    return doc


# ── Excel reading ─────────────────────────────────────────────────────────────
def normalise(s): return str(s).strip().lower() if s else ""

def read_sheet(ws, col_map):
    hri = None
    for row in ws.iter_rows():
        for cell in row:
            if normalise(cell.value) in col_map:
                hri = cell.row; break
        if hri: break
    if not hri: raise ValueError(f"Header row not found in '{ws.title}'")
    ci = {c.column: col_map[normalise(c.value)] for c in ws[hri] if normalise(c.value) in col_map}
    records = []
    for row in ws.iter_rows(min_row=hri+1):
        if all(c.value is None or str(c.value).strip()=='' for c in row): continue
        rec = {ci[c.column]: str(c.value).strip() if c.value is not None else '' for c in row if c.column in ci}
        if rec: records.append(rec)
    return records

def load_data(xp):
    wb = load_workbook(xp, data_only=True)
    for s in ("WI_Header","WI_Steps"):
        if s not in wb.sheetnames: sys.exit(f"ERROR: Sheet '{s}' not found in {xp}")
    return read_sheet(wb["WI_Header"], HEADER_COLS), read_sheet(wb["WI_Steps"], STEPS_COLS)

def find_excel(override=None):
    if override:
        p = Path(override)
        if p.exists(): return p
        p2 = SCRIPT_DIR / override
        if p2.exists(): return p2
        sys.exit(f"ERROR: Excel file not found: {override}")
    for name in EXCEL_FALLBACKS:
        p = SCRIPT_DIR / name
        if p.exists(): return p
    sys.exit(f"ERROR: Could not find Excel file in {SCRIPT_DIR}")


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input",    default=None)
    parser.add_argument("--template", default=None)
    parser.add_argument("--output",   default=None)
    parser.add_argument("--wi",       nargs="*")
    args = parser.parse_args()

    excel_path    = find_excel(args.input)
    template_path = Path(args.template) if args.template else DEFAULT_TEMPLATE
    output_dir    = Path(args.output)   if args.output   else DEFAULT_OUTPUT
    output_dir.mkdir(parents=True, exist_ok=True)

    if not template_path.exists():
        sys.exit(f"ERROR: Template not found: {template_path}")

    print(f"\n{'─'*60}\n  Work Instruction Document Generator\n{'─'*60}")
    print(f"  Excel    : {excel_path.name}")
    print(f"  Template : {template_path.name}")
    print(f"  Output   : {output_dir.resolve()}\n{'─'*60}\n")

    try:
        headers, steps = load_data(excel_path)
    except Exception as e:
        sys.exit(f"\nERROR reading Excel: {e}\nMake sure the file is closed.")

    print(f"  Loaded {len(headers)} WI(s), {len(steps)} step(s).")
    gen_vals = [h.get("Generate","(missing)") for h in headers]
    print(f"  Generate? values: {gen_vals}\n")

    to_gen = [h for h in headers if normalise(h.get("Generate","")).startswith("y")]
    if args.wi:
        wf = {w.upper() for w in args.wi}
        to_gen = [h for h in to_gen if h.get("WI_No","").upper() in wf]

    if not to_gen:
        print("  Nothing to generate. Set 'Generate?' to YES in the WI_Header sheet.\n")
        return

    print(f"  Generating {len(to_gen)} document(s)...\n")
    success = 0
    for wi in to_gen:
        wi_no = wi.get("WI_No","unknown")
        wi_steps = [s for s in steps if s.get("WI_No","") == wi_no]
        safe = wi_no.replace('/','').replace('\\','')
        out  = output_dir / f"{safe}.docx"
        print(f"  [{wi_no}]  {wi.get('Document_Title','')[:55]}")
        print(f"           {len(wi_steps)} steps  →  ", end="", flush=True)
        try:
            build_document(template_path, wi, wi_steps).save(str(out))
            print(f"✓  {out.name}"); success += 1
        except Exception as e:
            import traceback
            print(f"✗  FAILED — {e}")
            traceback.print_exc()

    print(f"\n{'─'*60}")
    print(f"  Done — {success}/{len(to_gen)} documents generated")
    print(f"  Saved to: {output_dir.resolve()}\n{'─'*60}\n")

if __name__ == "__main__":
    main()
